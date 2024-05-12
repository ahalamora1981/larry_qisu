import os
import shutil
import threading
import customtkinter
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


CWD = os.getcwd()

# 设定表格文件名，律师人名
EXCEL_FILE = 'info.xlsx'
LAWYER_LIST = ['王磊', '张立人', '杨青', '梅世正']

# 设定融担公司
RONGDAN_LIST = [
    '福建智云',
    '海南申信',
    '上海耳序'
]

# 合同号与律师对应关系
contract_to_lawyer = {}

# 证据文档列表
EVIDENCE_LIST = [
    '贷款合同',
    '个人委托担保协议',
    '不可撤销担保函',
    '借款服务协议',
    '债权转让协议'
]

INFO_COLUMNS = ['公司名称', '手别', '案件id', '用户ID', '用户名', '用户姓名', '身份证号码', '性别', '民族', '身份证地址', '注册手机号', '列表ID', '合同号', '资方机构', '融担公司', '合同金额', '放款日期', '最后一期应还款日', '借款期数', '利率', '逾期开始日期', '上一个还款日期', '列表逾期天数', '待还金额', '待还本金', '待还费用', '数据提取日', '合并 代偿回购总额', '合并 代偿回购本金', '合并 代偿回购利息', '合并 代偿回购罚息', '是否可诉', '最晚代偿回购时间', '管辖法院', '承办律师']

###############
### 文档处理 ###
###############
def process_docs_func(input_path, output_path):
    
    # 创建输出目录
    if os.path.exists(output_path):
        return "请删除 output 文件夹"
    else:
        os.mkdir(output_path)

    qsz_folder = os.path.join(input_path, '起诉状')
    wts_folder = os.path.join(input_path, '委托书')
    info_path = os.path.join(input_path, EXCEL_FILE)

    qsz_output_shex = os.path.join(output_path, '起诉状_上海耳序')
    qsz_output_fjzy = os.path.join(output_path, '起诉状_福建智云')
    qsz_output_hnsx = os.path.join(output_path, '起诉状_海南申信')

    wts_output_shex = os.path.join(output_path, '委托书_上海耳序')
    wts_output_fjzy = os.path.join(output_path, '委托书_福建智云')
    wts_output_hnsx = os.path.join(output_path, '委托书_海南申信')

    phone_number = '15900621166'

    # phone_wanglei = phone
    # phone_zhangliren = phone
    # phone_yangqing = phone

    phones = [
        '18916935832', # 王磊
        '13817213203', # 张立人
        '15221111951', # 杨青
        '15900621166' # 梅世正
    ]

    qsz_file_list = os.listdir(qsz_folder)
    wts_file_list = os.listdir(wts_folder)

    if not os.path.exists(os.path.join(output_path)):
        os.mkdir(os.path.join(output_path))

    if not os.path.exists(qsz_output_shex):
        os.mkdir(qsz_output_shex)
        
    if not os.path.exists(qsz_output_fjzy):
        os.mkdir(qsz_output_fjzy)
        
    if not os.path.exists(qsz_output_hnsx):
        os.mkdir(qsz_output_hnsx)

    if not os.path.exists(wts_output_shex):
        os.mkdir(wts_output_shex)
        
    if not os.path.exists(wts_output_fjzy):
        os.mkdir(wts_output_fjzy)
        
    if not os.path.exists(wts_output_hnsx):
        os.mkdir(wts_output_hnsx)

    df = pd.read_excel(info_path)

    num_shex = 0
    num_fjzy = 0
    num_hnsx = 0

    # 改管辖法院，并存到对应融担公司的目录中
    for qsz_file in qsz_file_list:
        
        # 读取起诉状文件
        qsz_path = os.path.join(qsz_folder, qsz_file)
        
        # 加载Docx文件
        doc = Document(qsz_path)
        
        # 从文件名获取合同号
        contract_id = qsz_file.split('_')[1]
        
        # 从表格中找到对应合同号的管辖法院
        court = df[df['合同号']==contract_id]['管辖法院'].tolist()[0]

        for p in doc.paragraphs:
            if '人民法院' in p.text:

                # 将正确的管辖法院更新到Docx中（加run设置字体）
                p.text = ''
                run = p.add_run(court)
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

                # 调整字体大小
                font = p.style.font
                font.size = Pt(14)
        
        # 判断担保公司
        company_fullname = df[df['合同号']==contract_id]['融担公司'].tolist()[0]
        
        if '福建智云' in company_fullname:
            output_folder = qsz_output_fjzy
            company = '福建智云'
            num_fjzy += 1
        elif '海南申信' in company_fullname:
            output_folder = qsz_output_hnsx
            company = '海南申信'
            num_hnsx += 1
        else:
            output_folder = qsz_output_shex
            company = '上海耳序'
            num_shex += 1
        
        person = df[df['合同号']==contract_id]['用户姓名'].tolist()[0]

        file_name = f'{person}起诉状{contract_id[-4:]}-拓棱特-诉状-{company}-电子.docx'
        
        # 保存文件到对应目录
        output_file_path = os.path.join(output_folder, file_name)
        
        if os.path.exists(output_file_path):
            return f"{output_file_path} 已存在！"
        else:
            doc.save(output_file_path)

    # 改律师名字和电话，并存到对应融担公司的目录中
    for wts_file in wts_file_list:
        
        # 读取委托书文件
        wts_path = os.path.join(wts_folder, wts_file)
        
        # 加载Docx文件
        doc = Document(wts_path)
        
        # 从文件名获取合同号
        contract_id = wts_file.split('_')[1]
        lawyer = df[df['合同号']==contract_id]['承办律师'].tolist()[0]
        
        for p in doc.paragraphs:
            if '王磊' in p.text:
                text_new_name = p.text.replace('王磊', lawyer)
                p.text = ''
                # 加run用于修改字体
                run = p.add_run(text_new_name)
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
                # 调整字体
                font = p.style.font
                font.size = Pt(14)
            
            for phone in phones:
                if phone in p.text:
                    text_new_phone = p.text.replace(phone, phone_number)
                    p.text = ''
                    # 加run用于修改字体
                    run = p.add_run(text_new_phone)
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
                    # 调整字体
                    font = p.style.font
                    font.size = Pt(14)
                    
                    break
        
        # 判断担保公司
        company_fullname = df[df['合同号']==contract_id]['融担公司'].tolist()[0]
        
        if '福建智云' in company_fullname:
            output_folder = wts_output_fjzy
            company = '福建智云'
        elif '海南申信' in company_fullname:
            output_folder = wts_output_hnsx
            company = '海南申信'
        else:
            output_folder = wts_output_shex
            company = '上海耳序'
            
        person = df[df['合同号']==contract_id]['用户姓名'].tolist()[0]

        file_name = f'{person}委托书{contract_id[-4:]}-拓棱特-委托书-{company}-电子.docx'
            
        # 保存文件到对应目录
        output_path = os.path.join(output_folder, file_name)
        
        if os.path.exists(output_path):
            raise FileExistsError
        else:
            doc.save(output_path)
            
    
    total_have = len(qsz_file_list)
    total_done = num_shex+num_fjzy+num_hnsx
    result = f'共 {total_have} 条\n完成 {total_done} 条\n\n上海耳序：共 {num_shex} 条 | 福建智云：共 {num_fjzy} 条 | 海南申信：共 {num_hnsx} 条\n\n所有文档已完成自动编辑！'
    
    return result

def process_docs():
    threading.Thread(target=process_docs_thread).start()

def process_docs_thread():
    input_path = os.path.join(CWD, '文档处理', 'input')
    output_path = os.path.join(CWD, '文档处理', 'output')
    
    if 'info.xlsx' not in os.listdir(input_path):
        raise FileNotFoundError('没有找到 info.xlsx 文件！')
    
    info_path = os.path.join(input_path, 'info.xlsx')
    
    info_check_result = check_info(info_path)
    
    if info_check_result is not True:
        app.tabs.tab1_textbox_check_info.delete("0.0", "end")
        app.tabs.tab1_textbox_check_info.insert("0.0", info_check_result)
    else:
        app.tabs.tab1_textbox_check_info.delete("0.0", "end")
        app.tabs.tab1_textbox_check_info.insert("0.0", "Info.xlsx 格式检查通过！")
        
        result = process_docs_func(input_path, output_path)
        app.tabs.tab1_textbox_result.insert("end", result + "\n")

###############
### 文档归类 ###
###############
def pack_docs_func(input_path, output_path):
    # 读取表格到 Pandas Dataframe
    df = pd.read_excel(os.path.join(input_path, EXCEL_FILE))

    # 创建输出目录
    if os.path.exists(output_path):
        return "请删除 output 文件夹"
    else:
        os.mkdir(output_path)

    # 创建律师目录
    for lawyer in LAWYER_LIST:
        if lawyer not in os.listdir(output_path):
            os.mkdir(os.path.join(output_path, lawyer))

    # 创建每个案件的目录
    for idx, row in enumerate(df[df['是否可诉']=='诉讼'].iterrows()):

        idx = '0' + str(idx + 1) if len(str(idx + 1)) == 1 else str(idx + 1)
        row = row[1]

        for rongdan in RONGDAN_LIST:
            if rongdan in row['融担公司']:
                break

        person = row['用户姓名']

        court = row['管辖法院'].strip('人民法院')

        contract_id = row['合同号']

        contract_to_lawyer[contract_id] = row['承办律师']

        case_folder = f'{idx}{rongdan}_{person} {court}-{contract_id}'
        
        if case_folder not in os.listdir(os.path.join(output_path, row['承办律师'])):
            os.mkdir(os.path.join(output_path, row['承办律师'], case_folder))

        if '委托材料' not in os.listdir(os.path.join(output_path, row['承办律师'], case_folder)):
            os.mkdir(os.path.join(output_path, row['承办律师'], case_folder, '委托材料'))

        if '证据' not in os.listdir(os.path.join(output_path, row['承办律师'], case_folder)):
            os.mkdir(os.path.join(output_path, row['承办律师'], case_folder, '证据'))
            
        if '债转通知' not in os.listdir(os.path.join(output_path, row['承办律师'], case_folder)):
            os.mkdir(os.path.join(output_path, row['承办律师'], case_folder, '债转通知'))
            
    result = "目录已创建"

    # 分发凭证
    pz_list = os.listdir(os.path.join(input_path, '凭证'))

    for pz in pz_list:
        list_id = int(pz.split('-')[0])
        
        if list_id in df['列表ID'].tolist():
            contract_id = df[df['列表ID']==list_id]['合同号'].tolist()[0]

            lawyer = contract_to_lawyer[contract_id]
            case_list = os.listdir(os.path.join(output_path, lawyer))
            case_folder = [folder for folder in case_list if contract_id in folder][0]
            case_path = os.path.join(output_path, lawyer, case_folder)

            if pz.split('-')[-1] not in os.listdir(os.path.join(case_path, '证据')):
                shutil.copy(
                    os.path.join(input_path, '凭证', pz), 
                    os.path.join(case_path, '证据', pz.split('-')[-1])
                )
        else:
            print(f"凭证: {list_id} 不存在")
            
    result += "\n凭证已分发"
    
    # 分发债转通知
    if '债转通知' in os.listdir(input_path):
        zztz_list = os.listdir(os.path.join(input_path, '债转通知'))

        for zztz in zztz_list:
            list_id = int(zztz.split('_')[1].split('.')[0])
            
            if list_id in df['列表ID'].tolist():
                contract_id = df[df['列表ID']==list_id]['合同号'].tolist()[0]
            
                lawyer = contract_to_lawyer[contract_id]
                case_list = os.listdir(os.path.join(output_path, lawyer))
                case_folder = [folder for folder in case_list if contract_id in folder][0]
                case_path = os.path.join(output_path, lawyer, case_folder)

                shutil.copy(
                    os.path.join(input_path, '债转通知', zztz), 
                    os.path.join(case_path, '债转通知', zztz)
                )
            else:
                print(f"债转通知: {list_id} 不存在")
        
        result += "\n债转通知已分发"

    # 分发文档包到证据目录
    doc_pack_level_2 = os.listdir(os.path.join(input_path, '文档包'))[0]
    doc_pack_list = os.listdir(os.path.join(input_path, '文档包', doc_pack_level_2))

    for doc_pack_folder in doc_pack_list:
        # 确定案件的输出目录路径 - 合同号 => 律师 => 案件输出路径
        contract_id = doc_pack_folder
        lawyer = contract_to_lawyer[contract_id]
        case_list = os.listdir(os.path.join(output_path, lawyer))
        case_folder = [folder for folder in case_list if contract_id in folder][0]
        case_path = os.path.join(output_path, lawyer, case_folder)

        # 遍历每个文件，再遍历每个符合要求的证据名称，确认文件是否为证据，是的话就复制到案件输出目录
        for file in os.listdir(os.path.join(input_path, '文档包', doc_pack_level_2, doc_pack_folder)):
            for evidence in EVIDENCE_LIST:
                if evidence in file:
                    if evidence + '.pdf' not in os.listdir(os.path.join(case_path, '证据')):
                        shutil.copy(
                            os.path.join(input_path, '文档包', doc_pack_level_2, doc_pack_folder, file),
                            os.path.join(case_path, '证据', evidence + '.pdf')
                        )
                        # print(os.path.join(case_path, '证据', evidence + '.pdf'))
                        
    result += "\n文档包已分发"

    # 分发文件到委托材料目录
    for lawyer in LAWYER_LIST:
        case_list = os.listdir(os.path.join(output_path, lawyer))

        for case_folder in case_list:
            # 确定融担公司
            for rongdan in RONGDAN_LIST:
                if rongdan in case_folder:
                    break
            
            # 确定合同号
            contract_id = case_folder.split('-')[-1]

            # 复制标准文档到委托材料
            standard_doc_list = os.listdir(os.path.join(input_path, "标准文档", rongdan))
            for standard_doc in standard_doc_list:
                if standard_doc not in os.listdir(os.path.join(output_path, lawyer, case_folder, '委托材料')):
                    shutil.copy(
                        os.path.join(input_path, "标准文档", rongdan, standard_doc),
                        os.path.join(output_path, lawyer, case_folder, '委托材料')
                    )

            # 复制委托书到委托材料
            wts_list = os.listdir(os.path.join(input_path, '委托书'))

            for wts in wts_list:
                name_wts = wts.split('_')[0]
                number_wts = wts.split('_')[1]
                if name_wts in case_folder and number_wts in case_folder:
                    if '授权委托书.pdf' not in os.listdir(os.path.join(output_path, lawyer, case_folder, '委托材料')):
                        shutil.copy(
                            os.path.join(input_path, '委托书', wts),
                            os.path.join(output_path, lawyer, case_folder, '委托材料', '授权委托书.pdf')
                        )

            # 复制身份证到委托材料
            username = df[df['合同号'] == contract_id]['用户名'].tolist()[0]

            id_list = os.listdir(os.path.join(input_path, '身份证_pdf'))

            for id_name in id_list:
                if username in id_name:
                    if '被告-身份证.pdf' not in os.listdir(os.path.join(output_path, lawyer, case_folder)):
                        shutil.copy(
                            os.path.join(input_path, '身份证_pdf', id_name),
                            os.path.join(output_path, lawyer, case_folder, '被告-身份证.pdf')
                        )

            # 复制起诉状到委托材料
            qsz_list = os.listdir(os.path.join(input_path, '起诉状'))

            for qsz in qsz_list:
                name_qsz = qsz.split('_')[0]
                number_qsz = qsz.split('_')[1]
                if name_qsz in case_folder and number_qsz in case_folder:
                    if '起诉状.pdf' not in os.listdir(os.path.join(output_path, lawyer, case_folder)):
                        shutil.copy(
                            os.path.join(input_path, '起诉状', qsz),
                            os.path.join(output_path, lawyer, case_folder, '起诉状.pdf')
                        )
                        
    result += "\n文件已分发"

    # # 复制一个新的输出目录，按律师要求把案件目录简化(去掉最后的合同号)
    # if 'output_final' not in os.listdir(base_path):
    #     shutil.copytree(output_path, os.path.join(base_path, 'output_final'))

    lawyer_list = os.listdir(output_path)

    for lawyer in lawyer_list:
        case_list = os.listdir(os.path.join(output_path, lawyer))

        for case_folder in case_list:
            if '-' in case_folder:
                os.rename(
                    os.path.join(output_path, lawyer, case_folder),
                    os.path.join(output_path, lawyer, case_folder.split('-')[0]),
                )
    
    return f"{result}\n\n全部完成!"

def pack_docs():
    threading.Thread(target=pack_docs_thread).start()

def pack_docs_thread():
    input_path = os.path.join(CWD, '文档归类', 'input')
    output_path = os.path.join(CWD, '文档归类', 'output')
    
    if 'info.xlsx' not in os.listdir(input_path):
        raise FileNotFoundError('没有找到 info.xlsx 文件！')
    
    info_path = os.path.join(input_path, EXCEL_FILE)
    info_check_result = check_info(info_path)
    
    if info_check_result is not True:
        app.tabs.tab2_textbox_check_info.delete("0.0", "end")
        app.tabs.tab2_textbox_check_info.insert("0.0", info_check_result)
    else:
        app.tabs.tab2_textbox_check_info.delete("0.0", "end")
        app.tabs.tab2_textbox_check_info.insert("0.0", "Info.xlsx 格式检查通过！")
        
        result = pack_docs_func(input_path, output_path)
        app.tabs.tab2_textbox_result.insert("end", result + "\n")

#############
### 律所函 ###
#############
def generate_mail_func(input_path, output_path): 
    # 创建输出目录
    if os.path.exists(output_path):
        return "请删除 output 文件夹"
    else:
        os.mkdir(output_path)
    
    info_path = os.path.join(input_path, EXCEL_FILE)
    
    df = pd.read_excel(info_path)
    df = df[df['是否可诉']=='诉讼'].reset_index()
    
    result = ""
    num_rows = df.shape[0]
    print(num_rows)
            
    for index, row in df.iterrows():
        row = row.astype('str')
        company_fullname = row['融担公司']
        
        if '福建智云' in company_fullname:
            company = '福建智云融资担保有限责任公司'
        elif '海南申信' in company_fullname:
            company = '海南申信融资担保有限公司'
        else:
            company = '上海耳序信息技术有限公司'

        user = row['用户姓名']
        court = row['管辖法院']
        lawyer = row['承办律师']
        
        doc = Document(os.path.join(input_path, '安徽苏滁律师事务所函.docx'))
        
        for p in doc.paragraphs:
            for r in p.runs:
                r.text = r.text.replace('【用户姓名】', user)
                r.text = r.text.replace('【融担公司】', company)
                r.text = r.text.replace('【管辖法院】', court)
                r.text = r.text.replace('【承办律师】', lawyer)

        doc_output_path = os.path.join(output_path, f"{user}--{company}--{court}--{lawyer}.docx")
        doc.save(doc_output_path)
        
        if lawyer == 'nan':
            error_msg = f"无承办律师: ({user}--{company}--{court}--{lawyer})"
            result += f"{error_msg}\n"
            
        app.tabs.tab3_progressbar.set((index + 1) / num_rows)
    
    return result if result else "已完成，一切正常！"
   
def generate_mail():
    threading.Thread(target=generate_mail_thread).start()

def generate_mail_thread():
    input_path = os.path.join(CWD, '律所函', 'input')
    output_path = os.path.join(CWD, '律所函', 'output')
    
    if 'info.xlsx' not in os.listdir(input_path):
        raise FileNotFoundError('没有找到 info.xlsx 文件！')
    
    info_path = os.path.join(input_path, 'info.xlsx')
    
    info_check_result = check_info(info_path)
    
    if info_check_result is not True:
        app.tabs.tab3_textbox_check_info.delete("0.0", "end")
        app.tabs.tab3_textbox_check_info.insert("0.0", info_check_result)
    else:
        app.tabs.tab3_textbox_check_info.delete("0.0", "end")
        app.tabs.tab3_textbox_check_info.insert("0.0", "Info.xlsx 格式检查通过！")
        
        result = generate_mail_func(input_path, output_path)
        app.tabs.tab3_textbox_result.insert("end", result + "\n")

#####################
### 检查 INFO 格式 ###
#####################
def check_info(info_path: os.PathLike) -> str | bool:
    df = pd.read_excel(info_path)
    
    output = ""
    
    for col in INFO_COLUMNS:
        if col not in df.columns:
            output += f"缺少字段：“{col}”。\n"

    if set(df['是否可诉'].unique()) != {'诉讼', '否'}:
        output += "“是否可诉”的选项有误，应为“诉讼”或“否”。"
        
    return output if output else True


customtkinter.set_appearance_mode("System")  # Modes: "System" (standard), "Dark", "Light"
customtkinter.set_default_color_theme("blue")  # Themes: "blue" (standard), "green", "dark-blue"


class TabView(customtkinter.CTkTabview):
    def __init__(self, master):
        super().__init__(master)
        self.grid_columnconfigure(0, weight=1)
        self.grid(row=0, column=0, padx=20, pady=20, sticky="nsew")
        self.configure()
        
        # create tabs
        self.add("文档处理")
        self.tab("文档处理").grid_columnconfigure(0, weight=1)
        
        self.add("文档归类")
        self.tab("文档归类").grid_columnconfigure(0, weight=1)
        
        self.add("律所函")
        self.tab("律所函").grid_columnconfigure(0, weight=1)
        
        #######################
        ### Tab 1 - 文档处理 ###
        #######################
        self.tab1_label = customtkinter.CTkLabel(
            self.tab("文档处理"), 
            text="文档处理", 
            font=customtkinter.CTkFont(size=24, weight="bold")
        )
        self.tab1_label.grid(row=0, column=0, padx=10, pady=10)
        
        self.tab1_btn = customtkinter.CTkButton(
            self.tab("文档处理"), 
            text="开始文档处理", 
            font=customtkinter.CTkFont(size=16, weight="bold"), 
            command=process_docs
        )
        self.tab1_btn.grid(row=1, column=0, padx=10, pady=10, ipadx=10, ipady=5, sticky="ew")
        
        self.tab1_textbox_check_info = customtkinter.CTkTextbox(
            self.tab("文档处理"), 
            height=50, 
            corner_radius=0,
            font=customtkinter.CTkFont(size=16)
        )
        self.tab1_textbox_check_info.grid(row=2, column=0, padx=10, pady=10, sticky="ew")
        self.tab1_textbox_check_info.insert("0.0", "待检查")
        
        self.tab1_textbox_result = customtkinter.CTkTextbox(
            self.tab("文档处理"), 
            height=285, 
            corner_radius=0,
            font=customtkinter.CTkFont(size=16)
        )
        self.tab1_textbox_result.grid(row=3, column=0, padx=10, pady=10, sticky="ew")

        #######################
        ### Tab 2 - 文档归类 ###
        #######################
        self.tab2_label = customtkinter.CTkLabel(
            self.tab("文档归类"), 
            text="文档归类", 
            font=customtkinter.CTkFont(size=24, weight="bold")
        )
        self.tab2_label.grid(row=0, column=0, padx=10, pady=10)
        
        self.tab2_btn = customtkinter.CTkButton(
            self.tab("文档归类"), 
            text="开始文档归类", 
            font=customtkinter.CTkFont(size=16, weight="bold"), 
            command=pack_docs
        )
        self.tab2_btn.grid(row=1, column=0, padx=10, pady=10, ipadx=10, ipady=5, sticky="ew")
        
        self.tab2_textbox_check_info = customtkinter.CTkTextbox(
            self.tab("文档归类"), 
            height=50, 
            corner_radius=0,
            font=customtkinter.CTkFont(size=16)
        )
        self.tab2_textbox_check_info.grid(row=2, column=0, padx=10, pady=10, sticky="ew")
        self.tab2_textbox_check_info.insert("0.0", "待检查")
        
        self.tab2_textbox_result = customtkinter.CTkTextbox(
            self.tab("文档归类"), 
            height=285, 
            corner_radius=0,
            font=customtkinter.CTkFont(size=16)
        )
        self.tab2_textbox_result.grid(row=3, column=0, padx=10, pady=10, sticky="ew")

        #########################
        ### Tab 3 - 律所函生成 ###
        #########################
        self.tab3_label = customtkinter.CTkLabel(
            self.tab("律所函"), 
            text="律所函", 
            font=customtkinter.CTkFont(size=24, weight="bold")
        )
        self.tab3_label.grid(row=0, column=0, padx=10, pady=10)
        
        self.tab3_btn = customtkinter.CTkButton(
            self.tab("律所函"), 
            text="生成律所函", 
            font=customtkinter.CTkFont(size=16, weight="bold"), 
            command=generate_mail
        )
        self.tab3_btn.grid(row=1, column=0, padx=10, pady=10, ipadx=10, ipady=5, sticky="ew")
        
        self.tab3_progressbar = customtkinter.CTkProgressBar(self.tab("律所函"), orientation="horizontal", height=10)
        self.tab3_progressbar.grid(row=2, column=0, padx=10, pady=10, sticky="ew")
        self.tab3_progressbar.set(0)
     
        self.tab3_textbox_check_info = customtkinter.CTkTextbox(
            self.tab("律所函"), 
            height=50, corner_radius=0,
            font=customtkinter.CTkFont(size=16)
        )
        self.tab3_textbox_check_info.grid(row=3, column=0, padx=10, pady=10, sticky="ew")
        self.tab3_textbox_check_info.insert("0.0", "待检查")
        
        self.tab3_textbox_result = customtkinter.CTkTextbox(
            self.tab("律所函"), 
            height=285, 
            corner_radius=0,
            font=customtkinter.CTkFont(size=16)
        )
        self.tab3_textbox_result.grid(row=4, column=0, padx=10, pady=10, sticky="ew")


class App(customtkinter.CTk):
    def __init__(self):
        super().__init__()

        # configure window
        self.title("法诉文档管理系统")
        self.geometry(f"{800}x{600}")

        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(0, weight=1)
        
        self.tabs = TabView(self)


if __name__ == "__main__":
    app = App()
    app.mainloop()