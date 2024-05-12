import os
import pandas as pd
from tqdm import tqdm
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


def doc_process(input_path, output_path):

    qsz_folder = os.path.join(input_path, '起诉状')
    wts_folder = os.path.join(input_path, '委托书')
    info_path = os.path.join(input_path, 'info.xlsx')

    qsz_output_shex = os.path.join(output_path, '起诉状_上海耳序')
    qsz_output_fjzy = os.path.join(output_path, '起诉状_福建智云')
    qsz_output_hnsx = os.path.join(output_path, '起诉状_海南申信')

    wts_output_shex = os.path.join(output_path, '委托书_上海耳序')
    wts_output_fjzy = os.path.join(output_path, '委托书_福建智云')
    wts_output_hnsx = os.path.join(output_path, '委托书_海南申信')

    phone_number = '15900621166'

    # phone_wanglei = phone
    # phone_zhangliren = phone
    # phone_yangqing = phone

    phones = [
        '18916935832', # 王磊
        '13817213203', # 张立人
        '15221111951', # 杨青
        '15900621166' # 梅世正
    ]

    qsz_file_list = os.listdir(qsz_folder)
    wts_file_list = os.listdir(wts_folder)

    if not os.path.exists(os.path.join(output_path)):
        os.mkdir(os.path.join(output_path))

    if not os.path.exists(qsz_output_shex):
        os.mkdir(qsz_output_shex)
        
    if not os.path.exists(qsz_output_fjzy):
        os.mkdir(qsz_output_fjzy)
        
    if not os.path.exists(qsz_output_hnsx):
        os.mkdir(qsz_output_hnsx)

    if not os.path.exists(wts_output_shex):
        os.mkdir(wts_output_shex)
        
    if not os.path.exists(wts_output_fjzy):
        os.mkdir(wts_output_fjzy)
        
    if not os.path.exists(wts_output_hnsx):
        os.mkdir(wts_output_hnsx)

    df = pd.read_excel(info_path)

    num_shex = 0
    num_fjzy = 0
    num_hnsx = 0

    # 改管辖法院，并存到对应融担公司的目录中
    for qsz_file in qsz_file_list:
        
        # 读取起诉状文件
        qsz_path = os.path.join(qsz_folder, qsz_file)
        
        # 加载Docx文件
        doc = Document(qsz_path)
        
        # 从文件名获取合同号
        contract_id = qsz_file.split('_')[1]
        
        # 从表格中找到对应合同号的管辖法院
        court = df[df['合同号']==contract_id]['管辖法院'].tolist()[0]

        for p in doc.paragraphs:
            if '人民法院' in p.text:

                # 将正确的管辖法院更新到Docx中（加run设置字体）
                p.text = ''
                run = p.add_run(court)
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

                # 调整字体大小
                font = p.style.font
                font.size = Pt(14)
        
        # 判断担保公司
        company_fullname = df[df['合同号']==contract_id]['融担公司'].tolist()[0]
        
        if '福建智云' in company_fullname:
            output_folder = qsz_output_fjzy
            company = '福建智云'
            num_fjzy += 1
        elif '海南申信' in company_fullname:
            output_folder = qsz_output_hnsx
            company = '海南申信'
            num_hnsx += 1
        else:
            output_folder = qsz_output_shex
            company = '上海耳序'
            num_shex += 1
        
        person = df[df['合同号']==contract_id]['用户姓名'].tolist()[0]

        file_name = f'{person}起诉状{contract_id[-4:]}-拓棱特-诉状-{company}-电子.docx'
        
        # 保存文件到对应目录
        output_path = os.path.join(output_folder, file_name)
        
        if os.path.exists(output_path):
            raise FileExistsError
        else:
            doc.save(output_path)

    # 改律师名字和电话，并存到对应融担公司的目录中
    for wts_file in wts_file_list:
        
        # 读取委托书文件
        wts_path = os.path.join(wts_folder, wts_file)
        
        # 加载Docx文件
        doc = Document(wts_path)
        
        # 从文件名获取合同号
        contract_id = wts_file.split('_')[1]
        lawyer = df[df['合同号']==contract_id]['承办律师'].tolist()[0]
        
        for p in doc.paragraphs:
            if '王磊' in p.text:
                text_new_name = p.text.replace('王磊', lawyer)
                p.text = ''
                # 加run用于修改字体
                run = p.add_run(text_new_name)
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
                # 调整字体
                font = p.style.font
                font.size = Pt(14)
            
            for phone in phones:
                if phone in p.text:
                    text_new_phone = p.text.replace(phone, phone_number)
                    p.text = ''
                    # 加run用于修改字体
                    run = p.add_run(text_new_phone)
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
                    # 调整字体
                    font = p.style.font
                    font.size = Pt(14)
                    
                    break
        
        # 从表格中找到对应合同号的管辖法院
        # lawyer = df[df['合同号']==contract_id]['承办律师'].tolist()[0]
        # user = df[df['合同号']==contract_id]['用户姓名'].tolist()[0]

        # 替换律师和电话
        # if lawyer != '王磊':
        #     num_other_lawyer += 1
        #     for p in doc.paragraphs:
        #         if '王磊' in p.text and user not in p.text:
        #             text_new_name = p.text.replace('王磊', '张立人')
        #             p.text = ''
        #             # 加run用于修改字体
        #             run = p.add_run(text_new_name)
        #             run.font.name = "Arial"
        #             run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
        #             # 调整字体
        #             font = p.style.font
        #             font.size = Pt(14)
        #             
        #         if phone_wanglei in p.text:
        #             text_new_phone = p.text.replace(phone_wanglei, phone_zhangliren)
        #             p.text = ''
        #             # 加run用于修改字体
        #             run = p.add_run(text_new_phone)
        #             run.font.name = "Arial"
        #             run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
        #             # 调整字体
        #             font = p.style.font
        #             font.size = Pt(14)

        #         if '王磊' in p.text and user in p.text:
        #             # 现委托 王磊 在我单位与 沈黎宾 追偿权纠纷案件中，作为我单位的委托代理人，代理权限如下：
        #             p.text = ''
        #             run = p.add_run('现委托')
        #             run.font.name = "Arial"
        #             run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

        #             run = p.add_run(' 张立人 ')
        #             run.font.underline = True
        #             run.font.name = "Arial"
        #             run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

        #             run = p.add_run('在我单位与')
        #             run.font.name = "Arial"
        #             run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

        #             run = p.add_run(f' {user} ')
        #             run.font.underline = True
        #             run.font.name = "Arial"
        #             run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

        #             run = p.add_run('追偿权纠纷案件中，作为我单位的委托代理人，代理权限如下：')
        #             run.font.name = "Arial"
        #             run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')                

        # 判断担保公司
        company_fullname = df[df['合同号']==contract_id]['融担公司'].tolist()[0]
        
        if '福建智云' in company_fullname:
            output_folder = wts_output_fjzy
            company = '福建智云'
        elif '海南申信' in company_fullname:
            output_folder = wts_output_hnsx
            company = '海南申信'
        else:
            output_folder = wts_output_shex
            company = '上海耳序'
            
        person = df[df['合同号']==contract_id]['用户姓名'].tolist()[0]

        file_name = f'{person}委托书{contract_id[-4:]}-拓棱特-委托书-{company}-电子.docx'
            
        # 保存文件到对应目录
        output_path = os.path.join(output_folder, file_name)
        
        if os.path.exists(output_path):
            raise FileExistsError
        else:
            doc.save(output_path)
            
    
    total_have = len(qsz_file_list)
    total_done = num_shex+num_fjzy+num_hnsx
    result = f'共 {total_have} 条\n完成 {total_done} 条\n\n上海耳序：共 {num_shex} 条 | 福建智云：共 {num_fjzy} 条 | 海南申信：共 {num_hnsx} 条\n\n所有文档已完成自动编辑！'
    
    return result

def mail_generate(input_path: os.PathLike, output_path: os.PathLike) -> None:
    info_path = os.path.join(input_path, 'info.xlsx')
    
    if not os.path.exists(output_path):
        os.mkdir(output_path)
    
    df = pd.read_excel(info_path)
    if set(df['是否可诉'].unique()) == {'诉讼', '否'}:
        print('OK')
    
    df = df[df['是否可诉']=='诉讼']
    
    for _, row in df.iterrows():
        row = row.astype('str')
        company_fullname = row['融担公司']
        
        if '福建智云' in company_fullname:
            company = '福建智云融资担保有限责任公司'
        elif '海南申信' in company_fullname:
            company = '海南申信融资担保有限公司'
        else:
            company = '上海耳序信息技术有限公司'

        user = row['用户姓名']
        court = row['管辖法院']
        lawyer = row['承办律师']
        
        break
        
        # doc = Document(os.path.join(input_path, '安徽苏滁律师事务所函.docx'))
        
        # for p in doc.paragraphs:
        #     for r in p.runs:
        #         r.text = r.text.replace('【用户姓名】', user)
        #         r.text = r.text.replace('【融担公司】', company)
        #         r.text = r.text.replace('【管辖法院】', court)
        #         r.text = r.text.replace('【承办律师】', lawyer)

        # doc_output_path = os.path.join(output_path, f"{user}--{company}--{court}--{lawyer}.docx")
        # doc.save(doc_output_path)
        
        # if lawyer == 'nan':
        #     print(f"Error: No Lawyer ({user}--{company}--{court}--{lawyer})")

def check_info(input_path: os.PathLike) -> None:
    cols = ['公司名称', '手别', '案件id', '用户ID', '用户名', '用户姓名', '身份证号码', '性别', '民族', '身份证地址', '注册手机号', '列表ID', '合同号', '资方机构', '融担公司', '合同金额', '放款日期', '最后一期应还款日', '借款期数', '利率', '逾期开始日期', '上一个还款日期', '列表逾期天数', '待还金额', '待还本金', '待还费用', '数据提取日', '合并 代偿回购总额', '合并 代偿回购本金', '合并 代偿回购利息', '合并 代偿回购罚息', '是否可诉', '最晚代偿回购时间', '管辖法院', '承办律师']
    
    info_path = os.path.join(input_path, 'info.xlsx')
    df = pd.read_excel(info_path)
    
    for col in cols:
        if col not in df.columns:
            print(f"Column “{col}” is missing!")
            break
    else:
        print("Columns are OK!")
    
    if set(df['是否可诉'].unique()) == {'诉讼', '否'}:
        print('“是否可诉” is OK!')
        
        
if __name__ == '__main__':
    CWD = os.getcwd()
    
    # doc_process()
    
    # mail_generate(
    #     os.path.join(CWD, "Documents", "律所函"), 
    #     os.path.join(CWD, "Documents","律所函", "output")
    # )
    
    check_info(os.path.join(CWD, "Documents", "律所函"))
    
    