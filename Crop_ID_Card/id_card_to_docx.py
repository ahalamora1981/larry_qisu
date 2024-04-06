import os
import shutil
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

CWD = os.getcwd()
print(CWD)

# 定义输入路径、输出路径和临时路径
output_path = os.path.join(CWD, 'output_temp')
temp_path = os.path.join(CWD, 'temp')
output_docx_folder = os.path.join(CWD, '身份证_docx')

output_list = [img_name for img_name in os.listdir(output_path) if '.jpg' in img_name]
output_path_list = [os.path.join(output_path, img_name) for img_name in output_list]

for output_path in output_path_list:
    if '正面' in output_path:
        doc = Document()
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        image_front_p = doc.add_paragraph()
        image_front_r = image_front_p.add_run()
        try:
            image_front_r.add_picture(output_path, width=Inches(5))
        
            image_front_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            doc.add_paragraph()
            
            image_back_p = doc.add_paragraph()
            image_back_r = image_back_p.add_run()
            image_back_r.add_picture(output_path.replace('正面', '反面'), width=Inches(5))
            image_back_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            print(output_path)
        
        if "\\" in output_path:
            splitter = "\\"
        else:
            splitter = "/"
            
        doc_file_name = os.path.join(output_docx_folder, output_path.split(splitter)[-1].split('_')[1] + '_身份证正反面.docx')
        
        doc.save(doc_file_name)
        
shutil.make_archive(output_docx_folder, 'zip', output_docx_folder)