import os
import cv2
import time
import numpy as np
import shutil
import matplotlib.pyplot as plt

from modelscope.pipelines import pipeline
from modelscope.utils.constant import Tasks
from modelscope.outputs import OutputKeys
from modelscope.utils.cv.image_utils import draw_face_detection_result
from modelscope.preprocessors.image import LoadImage

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_mask(img_path):
    result = salient_detect(img_path)
    cv2.imwrite('./temp/mask.jpg',result[OutputKeys.MASKS])
    
    # 读取掩码图像，确保它与原始图像具有相同的尺寸
    mask_image = cv2.imread('./temp/mask.jpg', cv2.IMREAD_GRAYSCALE)
    mask_image = cv2.GaussianBlur(mask_image, (5, 5), 0)
    
    return mask_image

def get_black_bg_image(img, mask_img):
    original_image = img
    output_image = np.zeros_like(original_image)
    
    # 遍历原始图像的每个像素
    for y in range(original_image.shape[0]):
        for x in range(original_image.shape[1]):
            # 检查掩码图像中的对应位置是否非零
            if mask_img[y, x] != 0:
                # 如果掩码非零，将原始图像中的像素复制到输出图像中
                output_image[y, x] = original_image[y, x]
    
    return output_image

def get_card_range(image):  # 传入掩码图片
    
    # 转换为灰度图像
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

    # 创建二值化图像，将非黑色像素设置为白色
    _, binary = cv2.threshold(gray, 1, 255, cv2.THRESH_BINARY)

    # 查找非黑色区域的轮廓
    contours, _ = cv2.findContours(binary, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    if contours:
        # 找到最上、最下、最左和最右的点
        largest_contour = max(contours, key=cv2.contourArea)
        leftmost = tuple(largest_contour[largest_contour[:, :, 0].argmin()][0])
        rightmost = tuple(largest_contour[largest_contour[:, :, 0].argmax()][0])
        topmost = tuple(largest_contour[largest_contour[:, :, 1].argmin()][0])
        bottommost = tuple(largest_contour[largest_contour[:, :, 1].argmax()][0])
    
    return {'left':leftmost, 'right':rightmost, 'top':topmost, 'bottom':bottommost}

def range_is_card(image, card_range):
    height, width = image.shape[:2]
    
    area_img = height * width
    area_card_range = (card_range['right'][0] - card_range['left'][0]) * (card_range['bottom'][1] - card_range['top'][1])
    
    ratio = area_card_range / area_img
    
    # print(f'Area_img: {area_img}\nArea_card_range: {area_card_range}\nRatio: {(area_card_range/area_img):.2f}')
    
    if card_range['left'][0] >= width/2 or card_range['right'][0] <= width/2:
        range_x_outside = True
    else:
        range_x_outside = False
        
    if card_range['top'][1] >= height/2 or card_range['bottom'][1] <= height/2:
        range_y_outside = True
    else:
        range_y_outside = False
    
    # 用面积比小于0.25来判断轮廓是头像，而不是偏一侧的身份证
    if (range_x_outside or range_y_outside) and ratio < 0.25:
        return False
    else:
        return True
    
def get_cropped_image(image, card_range): # image 为白底图片, img_range 为左右上下四个点的字典
    
    leftmost = card_range['left']
    rightmost = card_range['right']
    topmost = card_range['top']
    bottommost = card_range['bottom']

    # 计算截取区域的坐标
    x, y, w, h = leftmost[0], topmost[1], rightmost[0] - leftmost[0], bottommost[1] - topmost[1]

    # 截取图像
    cropped_image = image[y:y+h, x:x+w]

    return cropped_image

def get_rotate_image_90(image):
    # 获取图像的高度和宽度
    height, width = image.shape[:2]

    # 如果高度大于宽度，进行逆时针旋转90度
    if height > width:
        image = cv2.rotate(image, cv2.ROTATE_90_COUNTERCLOCKWISE)
        
    return image

def get_rotate_image_180(image):
    image = cv2.rotate(image, cv2.ROTATE_180)
        
    return image

def back_is_upside_down(image):
    height, width = image.shape[:2]
    
    # 分割图片为左右两个半区
    left_half = image[:, :width // 2]
    right_half = image[:, width // 2:]

    # 计算左右两个半区的红色像素数量
    red_pixels_left = count_red_pixels(left_half)
    red_pixels_right = count_red_pixels(right_half)

    # 比较红色像素数量并输出结果
    if red_pixels_left < red_pixels_right:
        return True
    else:
        return False
    
# 定义一个函数来计算红色像素的数量
def count_red_pixels(image):
    # 在HSV颜色空间中定义红色范围
    lower_red = np.array([0, 100, 100])
    upper_red = np.array([10, 255, 255])

    # 转换图片到HSV颜色空间
    hsv = cv2.cvtColor(image, cv2.COLOR_BGR2HSV)

    # 使用掩码来获取红色像素
    mask = cv2.inRange(hsv, lower_red, upper_red)
    red_pixels = cv2.countNonZero(mask)

    return red_pixels

def get_resized_image(image):
    # 获取原始图像的高度、宽度和通道数
    original_height, original_width, _ = image.shape
    
    # 设置新的高度为400像素，根据纵横比计算新的宽度
    new_height = 400
    new_width = int(original_width * (new_height / original_height))
    
    # 使用cv2.resize函数调整图像大小
    resized_image = cv2.resize(image, (new_width, new_height))
    
    return resized_image

def crop_id_image():
    # 如果输出路径不存在，则创建它
    if output_path not in os.listdir('./'):
        os.mkdir(output_path)
        
    # 如果临时路径不存在，则创建它
    if temp_path not in os.listdir('./'):
        os.mkdir(temp_path)

    # 如果临时路径不存在，则创建它
    if output_docx_folder not in os.listdir('./'):
        os.mkdir(output_docx_folder)

    # 获取输入路径中所有以'.jpg'结尾的图像文件名列表
    img_list = [img_name for img_name in os.listdir(input_path) if '.jpg' in img_name]

    # 构建完整的图像文件路径列表
    img_path_list = [os.path.join(input_path, img_name) for img_name in img_list]

    start = time.time()

    for img_name, img_path in zip(img_list, img_path_list):
    
        try:    
            # 读取图片到image对象
            image = cv2.imread(img_path)

            # 获得掩码图片
            mask_image = get_mask(img_path)

            # 用掩码图片获得黑底身份证图片
            black_bg_image = get_black_bg_image(image, mask_image)

            # 获得边框范围
            card_range = get_card_range(black_bg_image)

            # 判断边框是否为证件（如边框不是证件，则大概率是人像）
            if range_is_card(image, card_range):
                # 通过边框范围截取白底身份证图片
                cropped_image = get_cropped_image(image, card_range)

                # 如果高比宽长，就逆时针旋转90度
                image = get_rotate_image_90(cropped_image)

            else:
                pass

            cv2.imwrite('temp/temp.jpg', image)

            height, width = image.shape[:2]

            # 判断是否需要上下180度翻转
            if '正面' in img_path:
                result = face_detection('temp/temp.jpg')
                if result['boxes'][0][0] < width/2:
                    print(f'{img_path} 翻转180度')
                    image = get_rotate_image_180(image)
            else:
                if back_is_upside_down(image):
                    print(f'{img_path} 翻转180度')
                    image = get_rotate_image_180(image)

            image = get_resized_image(image)

            # 输出图片
            cv2.imwrite(os.path.join(output_path, 'output_' + img_name), image)

        except Exception as e:
            print(f'!!!ERROR!!! {img_path}:\n{e}\n')

    duration = time.time() - start

    print(duration)

def generate_docx(output_path, output_docx_folder):
    # 如果临时路径不存在，则创建它
    if output_docx_folder not in os.listdir('./'):
        os.mkdir(output_docx_folder)

    img_output_list = [img_name for img_name in os.listdir(output_path) if '.jpg' in img_name]
    img_output_path_list = [os.path.join(output_path, img_name) for img_name in img_output_list]

    for output_path in img_output_path_list:
        if '正面' in output_path:
            doc = Document()
            
            doc.add_paragraph()
            doc.add_paragraph()
            
            image_front_p = doc.add_paragraph()
            image_front_r = image_front_p.add_run()
            image_front_r.add_picture(output_path, width=Inches(5))
            image_front_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            doc.add_paragraph()
            
            image_back_p = doc.add_paragraph()
            image_back_r = image_back_p.add_run()
            image_back_r.add_picture(output_path.replace('正面', '反面'), width=Inches(5))
            image_back_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc_file_name = os.path.join(output_docx_folder, output_path.split('/')[-1].split('_')[1] + '_身份证正反面.docx')
            
            doc.save(doc_file_name)

if __name__ == '__main__':
    # 定义输入路径、输出路径和临时路径
    input_path = '身份证'
    output_path = 'output'
    temp_path = 'temp'
    output_docx_folder = '身份证_docx'

    salient_detect = pipeline(Tasks.semantic_segmentation, model='damo/cv_u2net_salient-detection')
    face_detection = pipeline(task=Tasks.face_detection, model='damo/cv_resnet_facedetection_scrfd10gkps')

    # crop_id_image(
    #     input_path,
    #     output_path,
    #     temp_path
    # )

    generate_docx(
        output_path,
        output_docx_folder
    )